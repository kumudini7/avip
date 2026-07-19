from __future__ import annotations

import re
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
from docx import Document as WordDocument

try:
    import pytesseract
except Exception:  # pragma: no cover - optional dependency
    pytesseract = None


def ensure_directory(path: str | Path) -> Path:
    target = Path(path)
    target.mkdir(parents=True, exist_ok=True)
    return target


def save_file(directory: str | Path, filename: str, content: bytes) -> Path:
    target_dir = ensure_directory(directory)
    safe_name = Path(filename).name
    target_path = target_dir / safe_name
    target_path.write_bytes(content)
    return target_path


def extract_text_from_pdf(file_path: Path) -> str:
    text_chunks: list[str] = []

    try:
        with fitz.open(file_path) as pdf:
            for page in pdf:
                page_text = page.get_text("text")
                if page_text:
                    text_chunks.append(page_text)
    except Exception:
        text_chunks = []

    if text_chunks:
        return "\n".join(text_chunks).strip()

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted = page.extract_text() or ""
                if extracted:
                    text_chunks.append(extracted)
    except Exception:
        pass

    if text_chunks:
        return "\n".join(text_chunks).strip()

    if pytesseract is not None:
        return _ocr_pdf_placeholder(file_path)

    return ""


def _ocr_pdf_placeholder(file_path: Path) -> str:
    try:
        with fitz.open(file_path) as pdf:
            if not pdf.page_count:
                return ""
            first_page = pdf.load_page(0)
            pix = first_page.get_pixmap(matrix=fitz.Matrix(2, 2))
            png_path = file_path.with_suffix(".page1.png")
            pix.save(png_path)
            try:
                return pytesseract.image_to_string(str(png_path)).strip()
            finally:
                if png_path.exists():
                    png_path.unlink(missing_ok=True)
    except Exception:
        return ""


def extract_text_from_docx(file_path: Path) -> str:
    document = WordDocument(file_path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    tables: list[str] = []
    for table in document.tables:
        for row in table.rows:
            row_values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_values:
                tables.append(" | ".join(row_values))
    return "\n".join(paragraphs + tables).strip()


def extract_text_from_file(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    if suffix in {".doc", ".docx"}:
        return extract_text_from_docx(file_path)
    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="ignore")
    return ""


def normalize_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()
